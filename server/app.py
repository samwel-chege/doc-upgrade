import os
from flask import Flask, jsonify, request, redirect, url_for, make_response, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from flask_bcrypt import Bcrypt
from flask_restful import Api, Resource
from datetime import timedelta
from sqlalchemy import func
from datetime import datetime, timezone
from werkzeug.utils import secure_filename
import spacy
from spacy.symbols import nsubjpass, agent
import PyPDF2
import docx
from flask_cors import CORS
from docx import Document as DocxDocument
import io
import re


from models import db, Document, Suggestion, DocumentHistory, StopWord, Word

# Initialize the Flask app
app = Flask(__name__)

CORS(app) 


# Configure the app for SQLite
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///docupp.db'  # Using SQLite
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SECRET_KEY'] = 'your_secret_key'  # Change to your secret key

# Initialize extensions

db.init_app(app) 
migrate = Migrate(app, db)
bcrypt = Bcrypt(app)
api = Api(app)  # Initialize the Api for flask_restful

# Simple route to check if the server is running
@app.route('/')
def index():
    return jsonify({"message": "Docupp spinned up succesfully!"})


# EXTRACT FILE

def extract_text(filepath):
    _, file_extension = os.path.splitext(filepath)

    if file_extension.lower() == '.txt':
        with open(filepath, 'r') as file:
            return file.read()
    elif file_extension.lower() == '.pdf':
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ''
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text()
            return text
    elif file_extension.lower() == '.docx':
        doc = docx.Document(filepath)
        return '\n'.join([para.text for para in doc.paragraphs])
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


# IMPROVE DOCUMENTS.

def improve_document(original_content):
    # Load the spaCy NLP model
    nlp = spacy.load("en_core_web_sm")
    
    # Process the text through the NLP pipeline
    doc = nlp(original_content)
    
    # Improved content (initially the same as original content)
    improved_content = original_content
    
    # Simplify words
    suggestions = []
    for token in doc:
        if token.is_alpha and len(token.text) > 6: 
            suggestions.append({
                "suggestion_text": f"Try simplifying: {token.text}",
                "original_word": token.text,  # Track the complex word
                "suggested_word": token.lemma_,  # Example: use lemma or another simpler word
                "suggestion_type": "simplify_word"
            })

    #Detect passive voice
    for sentence in doc.sents:
      for token in sentence:
        # Check if the token is a verb and has a passive subject
        if token.pos_ == "VERB" and any(child.dep_ == "nsubjpass" for child in token.children):
            # Check for auxiliary verbs associated with the verb
            if any(child.dep_ == "aux" or child.dep_ == "auxpass" for child in token.children):
                suggestions.append({
                    "suggestion_text": f"Passive voice detected: '{sentence.text.strip()}'. Consider rephrasing.",
                    "original_word": sentence.text.strip(),
                    "suggestion_type": "passive_voice"
                })
                break  # Only need to detect one instance per sentence

    # Detect repetition
    word_counts = {}
    for token in doc:
        if token.is_alpha:
            word = token.lemma_.lower()
            word_counts[word] = word_counts.get(word, 0) + 1      


     # Add suggestions for repeated words
    for word, count in word_counts.items():
        if count > 2:  # Customize threshold for what counts as repetitive
            suggestions.append({
                "suggestion_text": f"Repetition detected: The word '{word}' appears {count} times.",
                "original_word": word,
                "suggestion_type": "repetition"
            })                   

    return improved_content, suggestions



# POST DOCUMENT

UPLOAD_FOLDER = 'assets'
ALLOWED_EXTENSIONS = {'txt', 'docx', 'pdf'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/documents/upload', methods=['POST'])
def upload_document():
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file selected'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        # Extract text from the uploaded document (based on the file type)
        original_content = extract_text(filepath)

        # Process the document with spaCy to improve it
        improved_content, suggestions = improve_document(original_content)

        # Create a new Document object and save it in the database
        new_document = Document(
            upload_date=datetime.now(timezone.utc),
            original_content=original_content,
            improved_content=improved_content
        )

        db.session.add(new_document)
        db.session.commit()

        # Add each suggestion to the suggestions table
        for suggestion_data in suggestions:
            new_suggestion = Suggestion(
                document_id=new_document.id,
                suggestion_text=suggestion_data['suggestion_text'],  # Only the suggestion text
                original_word=suggestion_data.get('original_word', None),  # Store the original word
                suggested_word=suggestion_data.get('suggested_word', None),  # Store the suggested word
                suggestion_type=suggestion_data.get('suggestion_type', None),  # Store suggestion type
                is_accepted=False  # Default to not accepted
            )
            db.session.add(new_suggestion)

        db.session.commit()

        return jsonify({
            'message': 'Document uploaded and processed successfully',
            'original_content': original_content,
            'improved_content': improved_content,
            'suggestions': suggestions
        }), 201

    return jsonify({'error': 'Invalid file format'}), 400


# FETCH DOCUMENT

@app.route('/documents/latest', methods=['GET'])
def get_latest_document():
    
    # Fetch the most recent document 
    latest_document = Document.query.order_by(Document.upload_date.desc()).first()
    
    if not latest_document:
        return jsonify({'error': 'No documents found'}), 404

    # Serialize the document data
    document_data = {
        'id': latest_document.id, 
        'original_content': latest_document.original_content,
        'improved_content': latest_document.improved_content,
        # 'upload_date': latest_document.upload_date.isoformat()  
    }

    return jsonify({'latest_document': document_data}), 200


# FETCH SUGGESTION

@app.route('/documents/<int:document_id>/suggestions', methods=['GET'])

def get_document_suggestions(document_id):
    # Check if the document exists
    document = Document.query.get(document_id)

    if not document:
        return jsonify({'error': 'Document not found'}), 404

    # Fetch all suggestions for the given document
    suggestions = Suggestion.query.filter_by(document_id=document_id).all()

    if not suggestions:
        return jsonify({'message': 'No suggestions found for this document'}), 404

    # Manually serialize suggestions using to_dict method
    suggestions_data = [suggestion.to_dict() for suggestion in suggestions]

    return jsonify({
        'document_id': document_id,
        'suggestions': suggestions_data
    }), 200


# ACCEPT SUGGESTION

@app.route('/suggestions/accept/<int:suggestion_id>', methods=['POST'])
def accept_suggestion(suggestion_id):
    # Fetch the suggestion
    suggestion = Suggestion.query.get(suggestion_id)
    if not suggestion:
        return jsonify({'error': 'Suggestion not found'}), 404

    # Fetch the related document
    document = Document.query.get(suggestion.document_id)
    if not document:
        return jsonify({'error': 'Document not found'}), 404

    improved_content = document.improved_content

    # Handle suggestion based on its type
    if suggestion.suggestion_type == 'simplify_word':
        # Replace the original word with the suggested word in the improved content
        if suggestion.original_word and suggestion.suggested_word:
            improved_content = improved_content.replace(suggestion.original_word, suggestion.suggested_word)
    elif suggestion.suggestion_type == 'passive_voice':
        if suggestion.original_word:
            # Convert passive voice to active voice
            improved_sentence = convert_to_active_voice(suggestion.original_word)
            improved_content = improved_content.replace(suggestion.original_word, improved_sentence)
    elif suggestion.suggestion_type == 'repetition':
        # Remove extra occurrences of the repeated word while keeping one instance
        if suggestion.original_word:
            words = improved_content.split()
            filtered_words = []
            seen = set()
            for word in words:
                lower_word = word.lower()
                if lower_word == suggestion.original_word.lower():
                    if lower_word not in seen:
                        filtered_words.append(word)  # Keep the first occurrence
                        seen.add(lower_word)
                else:
                    filtered_words.append(word)
            improved_content = ' '.join(filtered_words)

    # Update the document's improved content and mark the suggestion as accepted
    document.improved_content = improved_content.strip()  # Strip to remove any extra spaces
    suggestion.is_accepted = True

    try:
        db.session.commit()
        return jsonify({
            'message': 'Suggestion accepted and applied successfully',
            'updated_improved_content': improved_content
        }), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Failed to accept suggestion: {str(e)}'}), 500
    
def convert_to_active_voice(passive_sentence):
    """
    Convert a passive voice sentence to active voice using spaCy.
    """
    nlp = spacy.load("en_core_web_sm")
    doc = nlp(passive_sentence)
    subject, agent, verb, aux = None, None, None, None

    for token in doc:
        if token.dep == nsubjpass:  # Passive subject
            subject = token.text
        elif token.dep == agent:  # Agent (introduced with "by")
            agent = " ".join([child.text for child in token.subtree])
        elif token.pos_ == "VERB":  # Main verb
            verb = token.lemma_
        elif token.dep_ == "aux":  # Auxiliary verb
            aux = token.text

    if subject and verb:
        # Build active sentence
        if agent:
            return f"{agent.replace('by ', '').strip().capitalize()} {verb} {subject}."
        else:
            return f"Someone {verb} {subject}."

    # Return original if conversion is not possible
    return passive_sentence    


# DENY SUGGESTION

@app.route('/suggestions/deny/<int:suggestion_id>', methods=['DELETE'])
def deny_suggestion(suggestion_id):
    # Get the suggestion
    suggestion = Suggestion.query.get(suggestion_id)

    if not suggestion:
        return jsonify({'error': 'Suggestion not found'}), 404

    # Delete the suggestion
    db.session.delete(suggestion)
    db.session.commit()

    return jsonify({'message': 'Suggestion denied and deleted'}), 200


# EXPORT DOCUMENT


@app.route('/export-document/<int:document_id>', methods=['GET'])
def export_document(document_id):
   
    
    # Fetch the document by ID
    document = Document.query.filter_by(id=document_id).first()

    if not document:
        return jsonify({'error': 'Document not found'}), 404

    # Create a new Word document using python-docx
    doc = DocxDocument()

    # Add content to the document (you can customize this)
    doc.add_heading('Improved Document', 0)
    doc.add_paragraph(document.improved_content)

    # Save the document to a BytesIO stream
    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)

    # Return the Word document as a downloadable file
    return send_file(
        output_stream,
        as_attachment=True,
        download_name="improved_document.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# Run the app
if __name__ == '__main__':
    app.run(port=5555)
